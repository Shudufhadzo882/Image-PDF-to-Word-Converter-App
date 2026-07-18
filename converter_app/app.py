import os
import io
import time
import uuid
import logging
import zipfile
from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Inches

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'pdf'}

# Attempt to configure Tesseract path for Windows users
tesseract_locations = [
    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    os.path.join(os.environ.get('USERPROFILE', ''), 'AppData', 'Local', 'Programs', 'Tesseract-OCR', 'tesseract.exe'),
]
tesseract_found = False
for loc in tesseract_locations:
    if os.path.exists(loc):
        pytesseract.pytesseract.tesseract_cmd = loc
        tesseract_found = True
        logger.info(f"Tesseract found at: {loc}")
        break

if not tesseract_found:
    logger.warning("Tesseract OCR executable not found in standard paths. System path fallback will be used.")

def is_tesseract_available():
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Fallback-aware PDF to images converter
def convert_pdf_to_images(pdf_bytes):
    # Try using PyMuPDF (fitz) first as it doesn't need poppler binaries
    try:
        import fitz
        images = []
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            # 150 DPI is standard for OCR
            pix = page.get_pixmap(dpi=150)
            img_data = pix.tobytes("png")
            images.append(Image.open(io.BytesIO(img_data)))
        pdf_doc.close()
        logger.info("PDF converted to images successfully using PyMuPDF (fitz).")
        return images
    except Exception as e:
        logger.warning(f"PyMuPDF conversion failed: {e}. Trying pdf2image...")
        try:
            from pdf2image import convert_from_bytes
            return convert_from_bytes(pdf_bytes)
        except Exception as inner_e:
            logger.error(f"pdf2image conversion also failed: {inner_e}")
            raise Exception(f"Failed to convert PDF pages to images. Please make sure PyMuPDF is installed or Poppler is configured. Details: {inner_e}")

def convert_image_to_docx(image_bytes, output_path):
    """Convert an image to a Word document with OCR text"""
    image = Image.open(io.BytesIO(image_bytes))
    doc = Document()
    
    # Extract text using OCR
    tesseract_ok = is_tesseract_available()
    if tesseract_ok:
        try:
            text = pytesseract.image_to_string(image)
            if text.strip():
                doc.add_paragraph(text)
        except Exception as e:
            logger.error(f"OCR error: {e}")
    else:
        logger.warning("Tesseract OCR not available. Skipping text extraction.")

    # Add the image to the document
    temp_image_name = f"temp_image_{uuid.uuid4().hex}.png"
    temp_image_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_image_name)
    image.save(temp_image_path)
    doc.add_picture(temp_image_path, width=Inches(6))
    
    try:
        os.remove(temp_image_path)
    except Exception as e:
        logger.error(f"Could not remove temp image: {e}")
    
    doc.save(output_path)
    return output_path

def convert_pdf_to_docx(pdf_bytes, output_path):
    """Convert PDF pages to images and extract text to Word document"""
    doc = Document()
    
    try:
        # Convert PDF pages to images
        images = convert_pdf_to_images(pdf_bytes)
        tesseract_ok = is_tesseract_available()
        
        for i, image in enumerate(images):
            # Extract text from each page using OCR
            if tesseract_ok:
                try:
                    text = pytesseract.image_to_string(image)
                    if text.strip():
                        doc.add_heading(f'Page {i + 1}', level=2)
                        doc.add_paragraph(text)
                except Exception as e:
                    logger.error(f"OCR error on page {i+1}: {e}")
            
            # Add page as image
            temp_image_name = f'temp_page_{i}_{uuid.uuid4().hex}.png'
            temp_image_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_image_name)
            image.save(temp_image_path, 'PNG')
            doc.add_picture(temp_image_path, width=Inches(6))
            doc.add_page_break()
            
            try:
                os.remove(temp_image_path)
            except Exception as e:
                logger.error(f"Could not remove temp image: {e}")
        
        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        # Fallback: create document with error message
        doc.add_paragraph(f"Error converting PDF: {str(e)}")
        doc.save(output_path)
        return output_path

@app.route('/')
def index():
    tesseract_ok = is_tesseract_available()
    return render_template('index.html', tesseract_ok=tesseract_ok)

@app.route('/convert', methods=['POST'])
def convert():
    if 'files' not in request.files:
        return jsonify({'error': 'No files uploaded'}), 400
    
    files = request.files.getlist('files')
    converted_files = []
    timestamp = str(int(time.time()))
    zip_filename = f'converted_{timestamp}.zip'
    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], zip_filename)
    
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_bytes = file.read()
                file_ext = filename.rsplit('.', 1)[1].lower()
                
                # Generate output filename
                base_name = filename.rsplit('.', 1)[0]
                output_filename = f"{base_name}.docx"
                output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{uuid.uuid4().hex}_{output_filename}")
                
                try:
                    if file_ext == 'pdf':
                        convert_pdf_to_docx(file_bytes, output_path)
                    else:
                        convert_image_to_docx(file_bytes, output_path)
                    
                    # Add to zip using original base name
                    zipf.write(output_path, output_filename)
                    converted_files.append(output_filename)
                    
                    # Clean up individual docx file
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    
                except Exception as e:
                    logger.error(f"Error converting {filename}: {e}")
                    if os.path.exists(output_path):
                        try: os.remove(output_path)
                        except: pass
                    continue
    
    if not converted_files:
        return jsonify({'error': 'No files could be converted'}), 400
    
    # Read the zip file into memory to avoid file lock during send_file on Windows
    with open(zip_path, 'rb') as f:
        zip_data = io.BytesIO(f.read())
    
    # Clean up zip file from output folder
    if os.path.exists(zip_path):
        try: os.remove(zip_path)
        except: pass
        
    return send_file(
        zip_data,
        as_attachment=True,
        download_name=f'converted_documents_{timestamp}.zip',
        mimetype='application/zip'
    )

@app.route('/convert_single', methods=['POST'])
def convert_single():
    """Convert a single file and return it directly"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if not file or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400
    
    filename = secure_filename(file.filename)
    file_bytes = file.read()
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    # Generate output filename
    base_name = filename.rsplit('.', 1)[0]
    output_filename = f"{base_name}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{uuid.uuid4().hex}_{output_filename}")
    
    try:
        if file_ext == 'pdf':
            convert_pdf_to_docx(file_bytes, output_path)
        else:
            convert_image_to_docx(file_bytes, output_path)
        
        # Read file into memory to avoid file lock during send_file on Windows
        with open(output_path, 'rb') as f:
            file_data = io.BytesIO(f.read())
            
        return send_file(
            file_data,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"Single file conversion failed: {e}")
        return jsonify({'error': f'Conversion failed: {str(e)}'}), 500
    finally:
        if os.path.exists(output_path):
            try: os.remove(output_path)
            except: pass

if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000)
